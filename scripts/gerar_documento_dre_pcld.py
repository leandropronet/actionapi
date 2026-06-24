#!/usr/bin/env python3
r"""Gera documento Word explicando, para o controller, a divergência de fórmula
de PCLD encontrada na planilha-modelo SGA_DRE E BP 2021 a 2025 ANUAL_v2.xlsx e
como o relatório DRE/BP gerado via ActionAPI (scripts/relatorio_dre.py) tratou
o caso.

Uso:

    .\.venv\Scripts\python.exe scripts\gerar_documento_dre_pcld.py
"""
from __future__ import annotations

import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor, Cm
except ModuleNotFoundError:
    print(
        "Dependência ausente. Execute:\n"
        "  .\\.venv\\Scripts\\python.exe -m pip install -r scripts\\requirements-documento-dre-pcld.txt",
        file=sys.stderr,
    )
    raise SystemExit(2)

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "relatorios" / "analise-divergencia-pcld-dre-vs-planilha-modelo.docx"

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
RED = RGBColor(0xB0, 0x20, 0x20)
GREEN = RGBColor(0x1E, 0x7A, 0x34)
GRAY = RGBColor(0x55, 0x55, 0x55)


def brl(value: float) -> str:
    text = f"{value:,.2f}"
    text = text.replace(",", "_").replace(".", ",").replace("_", ".")
    sign = "-" if value < 0 else ""
    return f"{sign}R$ {text.lstrip('-')}" if value < 0 else f"R$ {text}"


def add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
    return h


def add_para(doc: Document, text: str, bold: bool = False, italic: bool = False, color: RGBColor | None = None, size: int = 11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def style_table(table):
    table.style = "Light Grid Accent 1"
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    style_table(table)
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    return table


def build_document() -> Document:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    title = doc.add_heading("Análise de Divergência de Fórmula — PCLD na DRE", level=0)
    for run in title.runs:
        run.font.color.rgb = NAVY

    add_para(
        doc,
        f"Documento gerado em {datetime.now().strftime('%d/%m/%Y')} para validação do "
        "Controller. Compara a planilha-modelo SGA_DRE E BP 2021 a 2025 ANUAL_v2.xlsx "
        "com o relatório DRE/BP gerado diretamente via ActionAPI.",
        italic=True,
        color=GRAY,
    )

    # 1. Objetivo
    add_heading(doc, "1. Objetivo", level=1)
    add_para(
        doc,
        "Este documento descreve a metodologia usada para auditar as fórmulas da "
        "planilha-modelo de DRE/BP, os pontos de divergência encontrados e como o "
        "relatório DRE/BP gerado a partir da ActionAPI tratou cada ponto. O objetivo "
        "não é contestar o critério contábil definido pelo Controller — pelo contrário: "
        "a metodologia aplicada (tratar a Perda de PCLD como despesa real e as provisões "
        "de Constituição/Reversão como neutras na visão gerencial) foi mantida integralmente. "
        "O que foi corrigido é um erro de execução da fórmula em uma linha auxiliar da "
        "planilha, que fazia a Perda de PCLD ser descontada duas vezes do resultado contábil. "
        "Pedimos que o Controller valide se o raciocínio abaixo confere com a lógica que ele "
        "aplicou originalmente."
    )

    # 2. Período e fonte
    add_heading(doc, "2. Período e Fonte dos Dados", level=1)
    add_bullets(
        doc,
        [
            "Exercícios analisados: 2021 a 2025 (anos fechados). O relatório também pode "
            "ser gerado incluindo o ano corrente em andamento, como ano parcial.",
            "Fonte dos dados do relatório via API: tabela raw.contabil (réplica das tabelas "
            "contábeis do SIAGRI/Oracle — CABLANCTB + LANCONTAB — no PostgreSQL).",
            "Endpoint consultado: GET /api/v1/executivo/contabilidade/sintetico, com o "
            "parâmetro excluirEncerramento=true para a DRE (remove o histórico HIST_HIS = "
            "1000191, que zera as contas de resultado no encerramento anual). Para o Balanço "
            "Patrimonial não se usa esse filtro — o saldo acumulado normal é mantido.",
            "Cobertura de dados verificada diretamente no banco: o ano de 2021 está com os "
            "12 meses carregados (188.824 lançamentos), em volume comparável aos demais anos "
            "(180 mil em 2020, 202 mil em 2022) — não há lacuna de carga em 2021.",
        ],
    )

    # 3. Metodologia
    add_heading(doc, "3. Metodologia de Análise", level=1)
    add_para(
        doc,
        "Para não copiar cegamente uma fórmula da planilha-modelo que pudesse estar errada, "
        "a comparação seguiu estes passos:"
    )
    add_bullets(
        doc,
        [
            "Passo 1 — Leitura das fórmulas, não dos valores: a planilha-modelo foi aberta "
            "em modo de leitura das fórmulas originais (não dos valores já calculados pelo "
            "Excel), na aba \"SGA_DRE Comparativa Exercicio\", para examinar exatamente como "
            "cada linha é calculada, célula a célula.",
            "Passo 2 — Reconstrução equivalente no relatório da API: o script "
            "scripts/relatorio_dre.py replica a mesma estrutura de contas e fórmulas da "
            "planilha-modelo, com saldos vindos da ActionAPI em vez de cálculos manuais em "
            "Excel.",
            "Passo 3 — Expansão algébrica das fórmulas: cada fórmula foi reescrita em termos "
            "de suas contas de origem (substituição simbólica), para identificar se algum "
            "valor aparecia somado ou subtraído mais de uma vez ao longo da cadeia de cálculo.",
            "Passo 4 — Confronto numérico por ano: os valores de 2021 a 2025 de cada conta "
            "envolvida foram extraídos da própria planilha-modelo, para isolar em qual(is) "
            "exercício(s) a divergência teria efeito monetário real.",
            "Passo 5 — Verificação de impacto no resultado final: para cada divergência "
            "encontrada, foi verificado se ela afeta o \"Resultado do Exercício\" (linha que "
            "alimenta ROA, ROE e EBITDA) ou apenas uma linha auxiliar/informativa.",
        ],
    )

    # 4. Achados
    add_heading(doc, "4. Achados", level=1)

    add_heading(doc, "4.1 Rótulo duplicado nas linhas 59 e 60", level=2)
    add_para(
        doc,
        "Na aba \"SGA_DRE Comparativa Exercicio\", as linhas 59 e 60 têm o mesmo texto na "
        "coluna C: \"RESULTADO CONTÁBIL ANTES DOS IMPOSTOS\". São, porém, dois cálculos "
        "diferentes:"
    )
    add_table(
        doc,
        ["Linha", "Rótulo na planilha", "Fórmula (coluna E, ano 2025)", "O que de fato calcula"],
        [
            ["59", "RESULTADO CONTÁBIL ANTES DOS IMPOSTOS", "=-(E53-E57+E58)+E49+E50",
             "Resultado considerando o efeito completo do PCLD (Perda + Constituição + Reversão)"],
            ["60", "RESULTADO CONTÁBIL ANTES DOS IMPOSTOS (rótulo duplicado)", "=E49+E50+E57-E58",
             "Resultado gerencial: exclui o efeito das provisões de PCLD (Constituição/Reversão)"],
        ],
    )
    add_para(
        doc,
        "Entendimento: a linha 60 deveria se chamar \"RESULTADO GERENCIAL ANTES DOS "
        "IMPOSTOS\" — é um erro de digitação/cópia, não um erro de cálculo. A lógica de "
        "manter duas visões (uma com o efeito completo do PCLD, outra sem o efeito das "
        "provisões) está correta; só o texto da linha 60 foi copiado da linha 59 por engano.",
        color=GREEN,
    )

    add_heading(doc, "4.2 Perda de PCLD descontada em dobro na linha 59", level=2)
    add_para(
        doc,
        "Este é o achado com impacto monetário. A conta 4211250060 (Despesa com Perda de "
        "PCLD) é contabilizada dentro do grupo 4211 (Despesas Administrativas e Comerciais). "
        "A fórmula da linha de Despesas (linha 35) é:"
    )
    add_para(doc, "E35 = SOMA(grupo 4211) - E55 (Constituição) - E56 (Reversão) - E58 (Perdas de Estoque)", bold=True)
    add_para(
        doc,
        "Ou seja: das contas do grupo 4211, a planilha retira Constituição, Reversão e "
        "Perdas de Estoque — mas NÃO retira a Perda de PCLD (E54). A Perda de PCLD "
        "permanece dentro das Despesas e, portanto, já reduziu o Lucro Operacional (linha 49)."
    )
    add_para(
        doc,
        "A linha 53 (PCLD) soma os três componentes: E53 = E54 (Perda) + E55 (Constituição) "
        "+ E56 (Reversão). Quando a linha 59 desconta E53 inteiro do resultado, a Perda de "
        "PCLD (E54) é subtraída uma segunda vez — ela já tinha sido descontada dentro de "
        "E35/E49, e é descontada de novo via E53."
    )
    add_para(doc, "Demonstração algébrica (expandindo E59 em função das contas de origem):", bold=True)
    add_table(
        doc,
        ["Etapa", "Expressão"],
        [
            ["Fórmula original da linha 59", "E59 = -(E53 - E57 + E58) + E49 + E50"],
            ["Substituindo E53", "E59 = -(E54 + E55 + E56 - E57 + E58) + E49 + E50"],
            ["Substituindo E49 (já contém +E55+E56+E58 internamente, vindos do cancelamento de E35)",
             "E59 = -E54 + E57 + [Lucro Bruto - Despesas(grupo 4211 bruto) + Resultado Financeiro]"],
            ["Resultado", "A Perda de PCLD (E54) aparece com sinal negativo MESMO já estando embutida "
             "e descontada dentro de \"Despesas (grupo 4211 bruto)\" — dupla contagem."],
        ],
    )
    add_para(
        doc,
        "Impacto: o valor de E59 fica subestimado exatamente no valor da Perda de PCLD do "
        "ano. A fórmula correta da linha 59 deve descontar apenas Constituição + Reversão "
        "(não a Perda):"
    )
    add_para(doc, "Fórmula corrigida: E59 = -((E55 + E56) - E57 + E58) + E49 + E50", bold=True, color=RED)

    add_heading(doc, "4.3 Valor encontrado em 2021 e por que não aparece nos demais anos", level=2)
    add_para(
        doc,
        "A conta 4211250060 (Perda de PCLD) só teve saldo relevante em 2021. Nos anos "
        "seguintes ficou zerada — por isso o erro de fórmula é \"invisível\" ao olhar "
        "qualquer ano de 2022 a 2025: subtrair zero duas vezes ainda dá zero."
    )
    add_table(
        doc,
        ["Ano", "Perda de PCLD (conta 4211250060)", "Resultado Contábil — linha 59 (com erro)", "Resultado Contábil corrigido", "Diferença"],
        [
            ["2021", brl(2_287_014.42), brl(15_877_758.91), brl(18_164_773.33), brl(2_287_014.42)],
            ["2022", brl(0.0), brl(49_313_781.02), brl(49_313_781.02), brl(0.0)],
            ["2023", brl(0.0), brl(26_817_316.33), brl(26_817_316.33), brl(0.0)],
            ["2024", brl(0.0), brl(-4_107_129.25), brl(-4_107_129.25), brl(0.0)],
            ["2025", brl(0.0), brl(-12_378_943.52), brl(-12_378_943.52), brl(0.0)],
        ],
    )
    add_para(
        doc,
        "A diferença em cada ano é exatamente igual ao saldo da conta de Perda de PCLD "
        "daquele ano — confirmação adicional de que a causa do desvio é a dupla contagem "
        "dessa conta, e não outro fator.",
        italic=True,
        color=GRAY,
    )
    add_para(
        doc,
        "Importante para a validação do Controller: esse erro NÃO afeta a linha 60 "
        "(\"Resultado Gerencial\") nem a linha 63 (\"RESULTADO DO EXERCÍCIO\"), que é o "
        "número que efetivamente alimenta os indicadores de ROA, ROE e EBITDA. Ou seja, o "
        "lucro final reportado pela empresa nos demonstrativos não está e nunca esteve "
        "incorreto por esse motivo — o problema está restrito a uma linha auxiliar/"
        "informativa (linha 59) que mostra o resultado \"com efeito completo do PCLD\".",
        bold=True,
    )

    # 5. Tratamento no relatório da API
    add_heading(doc, "5. Como o relatório DRE/BP da ActionAPI tratou o caso", level=1)
    add_bullets(
        doc,
        [
            "Os dois rótulos foram mantidos distintos desde a primeira versão do relatório: "
            "\"Resultado Contábil antes dos Impostos (com PCLD)\" e \"Resultado Gerencial "
            "antes dos Impostos\" — eliminando a duplicidade de nome do item 4.1.",
            "A fórmula da linha \"Resultado Contábil antes dos Impostos (com PCLD)\" foi "
            "ajustada para descontar apenas Constituição + Reversão do PCLD, não a Perda "
            "(que já está embutida nas Despesas) — eliminando a dupla contagem do item 4.2.",
            "O \"Resultado do Exercício\" final do relatório da API segue exatamente a mesma "
            "metodologia do Controller: a Perda de PCLD permanece como despesa real dentro "
            "do resultado; as provisões de Constituição/Reversão são neutralizadas apenas na "
            "visão gerencial. Nenhuma decisão de critério contábil foi alterada — somente a "
            "aritmética de uma linha auxiliar.",
            "O relatório inclui uma aba \"Auditoria Fórmulas\" que documenta este e outros "
            "pontos de forma explícita (fórmula usada, fórmula observada no modelo, nível de "
            "risco e decisão tomada), para rastreabilidade.",
        ],
    )

    # 6. Outros pontos (menor risco)
    add_heading(doc, "6. Outros pontos observados (informativo, menor risco)", level=1)
    add_bullets(
        doc,
        [
            "ROA e ROE: o relatório da API usa o Resultado do Exercício da própria DRE "
            "gerada, em vez do saldo de Lucros/Prejuízos Acumulados do Balanço Patrimonial "
            "(que pode ficar zerado por distribuição ou encerramento e distorcer a série "
            "histórica).",
            "Liquidez Geral e Endividamento: a planilha-modelo usa critérios que diferem da "
            "definição técnica usual (ex.: inclui o Ativo Não Circulante inteiro em vez de "
            "apenas o Realizável a Longo Prazo). O relatório da API mostra os dois critérios "
            "lado a lado, sem decidir qual é o certo, para reconciliação com o Controller.",
        ],
    )

    # 7. Pedido de validação
    add_heading(doc, "7. Pontos para validação do Controller", level=1)
    add_bullets(
        doc,
        [
            "Confirmar que a metodologia de tratar a Perda de PCLD (conta 4211250060) como "
            "despesa real — mantida no resultado — e as provisões de Constituição/Reversão "
            "como neutras na visão gerencial reflete corretamente o critério contábil da "
            "empresa.",
            "Validar a correção proposta na fórmula da linha 59 da planilha-modelo "
            "(descontar apenas Constituição + Reversão, não a Perda).",
            "Validar a correção do rótulo da linha 60 para \"RESULTADO GERENCIAL ANTES DOS "
            "IMPOSTOS\".",
            "Indicar se deseja que essas duas correções sejam replicadas na planilha-modelo "
            "oficial (até o momento, a planilha original não foi alterada; as correções "
            "estão aplicadas apenas no script que gera o relatório via API).",
        ],
    )

    add_para(doc, "")
    add_para(
        doc,
        "Fontes consultadas: relatorios/contabilidade/SGA_DRE E BP 2021 a 2025 ANUAL_v2.xlsx "
        "(aba SGA_DRE Comparativa Exercicio) e scripts/relatorio_dre.py.",
        italic=True,
        size=9,
        color=GRAY,
    )
    return doc


def main() -> None:
    doc = build_document()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"[documento] arquivo: {OUTPUT}")


if __name__ == "__main__":
    main()
